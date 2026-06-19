import io
import base64
import hashlib
import json
from pathlib import Path
from dataclasses import dataclass
from typing import AsyncIterator

import tiktoken
from PIL import Image
import fitz  # PyMuPDF

from app.utils.config import IMAGE_MAX_DIMENSION, IMAGE_SAVE_DIR

SUPPORTED_EXTENSIONS = {".pdf", ".md", ".txt", ".docx"}
MAX_FILE_SIZE_BYTES = 10 * 1024 * 1024  # 10 MB

_tiktoken_enc = tiktoken.get_encoding("cl100k_base")


class ParseError(Exception):
    """Raised when document parsing fails."""
    pass


class UnsupportedFileTypeError(ParseError):
    """Raised when the file type is not supported."""
    pass


class FileTooLargeError(ParseError):
    """Raised when the file exceeds the size limit."""
    pass


class FileContentInvalidError(ParseError):
    """Raised when the file content doesn't match its claimed extension."""
    pass


# ---------------------------------------------------------------------------
# Dataclasses
# ---------------------------------------------------------------------------

@dataclass
class ImageItem:
    """An extracted image with its caption and metadata."""
    image_id: str
    doc_id: str
    page_num: int
    image_index: int          # index within the page
    file_path: str            # disk path to saved image
    caption: str              # Vision LLM generated caption
    width: int
    height: int
    image_base64: str = ""    # base64 encoded for Vision LLM prompts


@dataclass
class ParsedDocument:
    """Result of parsing a multimodal document."""
    doc_id: str
    filename: str
    file_type: str
    full_text: str            # all extracted text (from pdfplumber or fitz)
    text_chunks: list[dict]   # [{"content": str, "index": int}, ...]
    images: list[ImageItem]   # extracted images


# ---------------------------------------------------------------------------
# Validation
# ---------------------------------------------------------------------------

def validate_file(filename: str, content: bytes) -> Path:
    """Validate file type, size, and content integrity."""
    if not filename:
        raise UnsupportedFileTypeError("文件名为空")

    ext = Path(filename).suffix.lower()
    if not ext:
        raise UnsupportedFileTypeError("无法识别文件类型，缺少文件扩展名")
    if ext not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise UnsupportedFileTypeError(f"不支持的文件类型: {ext}。支持的类型: {supported}")
    if len(content) == 0:
        raise FileContentInvalidError("文件内容为空，请上传有效文件")
    if len(content) > MAX_FILE_SIZE_BYTES:
        max_mb = MAX_FILE_SIZE_BYTES // (1024 * 1024)
        raise FileTooLargeError(
            f"文件大小 ({len(content) / (1024*1024):.1f}MB) 超过限制 ({max_mb}MB)"
        )
    return Path(filename)


# ---------------------------------------------------------------------------
# PDF Parsing — Text + Image Extraction
# ---------------------------------------------------------------------------

def _compress_image(pil_img: Image.Image, max_dim: int = IMAGE_MAX_DIMENSION) -> Image.Image:
    """Resize image so its largest dimension <= max_dim, preserving aspect ratio."""
    w, h = pil_img.size
    if max(w, h) <= max_dim:
        return pil_img
    scale = max_dim / max(w, h)
    new_size = (int(w * scale), int(h * scale))
    return pil_img.resize(new_size, Image.LANCZOS)


def _image_to_base64(pil_img: Image.Image) -> str:
    """Convert PIL Image to base64 data URL string."""
    buf = io.BytesIO()
    # Convert RGBA to RGB for JPEG
    if pil_img.mode in ("RGBA", "P"):
        pil_img = pil_img.convert("RGB")
    pil_img.save(buf, format="JPEG", quality=85)
    return base64.b64encode(buf.getvalue()).decode("utf-8")


def _compute_image_hash(img_bytes: bytes) -> str:
    """Compute a short hash for image deduplication."""
    return hashlib.md5(img_bytes).hexdigest()[:12]


def _extract_images_from_pdf(
    doc_id: str,
    pdf_content: bytes,
    image_save_dir: str = IMAGE_SAVE_DIR,
) -> list[ImageItem]:
    """Extract embedded images from a PDF using PyMuPDF (fitz).

    Returns a list of ImageItem dataclasses with saved images and metadata.
    """
    images: list[ImageItem] = []
    doc_dir = Path(image_save_dir) / doc_id
    doc_dir.mkdir(parents=True, exist_ok=True)

    pdf_doc = fitz.open(stream=pdf_content, filetype="pdf")

    for page_num in range(len(pdf_doc)):
        page = pdf_doc[page_num]
        # Get all images on this page
        image_list = page.get_images(full=True)

        for img_index, img_info in enumerate(image_list):
            xref = img_info[0]
            try:
                base_image = pdf_doc.extract_image(xref)
                img_bytes = base_image["image"]
                ext = base_image["ext"]  # e.g., 'png', 'jpeg'
            except Exception:
                continue

            # Skip tiny images (likely icons / decorations)
            if len(img_bytes) < 1024:
                continue

            # Compute hash and save
            img_hash = _compute_image_hash(img_bytes)
            image_id = f"{doc_id}_p{page_num}_i{img_index}_{img_hash}"
            file_name = f"{image_id}.{ext}"
            file_path = str(doc_dir / file_name)

            # Open with PIL, compress if needed
            pil_img = Image.open(io.BytesIO(img_bytes))
            orig_w, orig_h = pil_img.size
            pil_img = _compress_image(pil_img)

            # Save to disk
            pil_img.save(file_path)

            images.append(ImageItem(
                image_id=image_id,
                doc_id=doc_id,
                page_num=page_num,
                image_index=img_index,
                file_path=file_path,
                caption="",  # filled later by Vision LLM
                width=orig_w,
                height=orig_h,
            ))

    pdf_doc.close()
    return images


def _extract_text_from_pdf(pdf_content: bytes) -> str:
    """Extract text from PDF using pdfplumber (primary) with PyMuPDF fallback."""
    # Primary: pdfplumber
    import pdfplumber

    text_parts: list[str] = []
    with pdfplumber.open(io.BytesIO(pdf_content)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

    if text_parts:
        return "\n\n".join(text_parts)

    # Fallback: PyMuPDF
    pdf_doc = fitz.open(stream=pdf_content, filetype="pdf")
    for page in pdf_doc:
        page_text = page.get_text()
        if page_text and page_text.strip():
            text_parts.append(page_text.strip())
    pdf_doc.close()

    if not text_parts:
        raise ParseError("PDF 文件中未提取到文本内容，可能是扫描件或图片型 PDF")

    return "\n\n".join(text_parts)


# ---------------------------------------------------------------------------
# Vision LLM — Image Captioning
# ---------------------------------------------------------------------------

CAPTION_PROMPT = """请用一段中文描述这张图片的内容，包括：
1. 图片类型（如架构图、流程图、数据图表、截图、照片等）
2. 图片中的关键视觉元素和文字信息
3. 图片传达的核心信息或结论

请控制在 150 字以内，用简洁的说明性语言。"""


async def generate_image_caption(
    image_base64: str,
    prompt: str = CAPTION_PROMPT,
) -> str:
    """Call Vision LLM to generate a caption for an image."""
    from app.services.llm import get_vision_client, get_vision_model

    client = get_vision_client()
    model = get_vision_model()

    response = await client.chat.completions.create(
        model=model,
        messages=[{
            "role": "user",
            "content": [
                {
                    "type": "image_url",
                    "image_url": {"url": f"data:image/jpeg;base64,{image_base64}"},
                },
                {"type": "text", "text": prompt},
            ],
        }],
        max_tokens=256,
        temperature=0.3,
    )
    return response.choices[0].message.content or ""


async def caption_images(images: list[ImageItem]) -> list[ImageItem]:
    """Generate captions for all images using Vision LLM.

    Processes images sequentially to avoid rate limiting.
    """
    for img in images:
        # Load and encode the saved image
        pil_img = Image.open(img.file_path)
        img.image_base64 = _image_to_base64(pil_img)
        pil_img.close()

        try:
            img.caption = await generate_image_caption(img.image_base64)
        except Exception as e:
            img.caption = f"[图片摘要生成失败: {str(e)[:100]}]"

    return images


# ---------------------------------------------------------------------------
# Plain-text Parsers
# ---------------------------------------------------------------------------

def parse_markdown(content: bytes) -> str:
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        return content.decode("utf-8", errors="replace")


def parse_txt(content: bytes) -> str:
    for encoding in ("utf-8", "gbk", "gb2312", "latin-1"):
        try:
            return content.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    return content.decode("utf-8", errors="replace")


def parse_docx(content: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(content))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    tables_text: list[str] = []

    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text for cell in row.cells if cell.text.strip()]
            if row_texts:
                tables_text.append(" | ".join(row_texts))

    if not paragraphs and not tables_text:
        raise ParseError("DOCX 文件中未提取到文本内容，文件可能为空")

    result_parts: list[str] = []
    if paragraphs:
        result_parts.append("\n\n".join(paragraphs))
    if tables_text:
        result_parts.append("\n\n--- 表格内容 ---\n" + "\n".join(tables_text))

    return "\n\n".join(result_parts)


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

async def parse_document(
    filename: str,
    content: bytes,
    doc_id: str,
) -> ParsedDocument:
    """Parse a document and extract text + images.

    For PDF files: text via pdfplumber, images via PyMuPDF, captions via Vision LLM.
    For other files: text-only extraction.

    Raises:
        UnsupportedFileTypeError, FileTooLargeError, FileContentInvalidError, ParseError
    """
    file_path = validate_file(filename, content)
    ext = file_path.suffix.lower()

    # Determine file type for metadata
    file_type = ext.lstrip(".")

    # Parse text
    if ext == ".pdf":
        full_text = _extract_text_from_pdf(content)
        images = _extract_images_from_pdf(doc_id, content)
    elif ext == ".md":
        full_text = parse_markdown(content)
        images = []
    elif ext == ".txt":
        full_text = parse_txt(content)
        images = []
    elif ext == ".docx":
        full_text = parse_docx(content)
        images = []
    else:
        raise UnsupportedFileTypeError(f"不支持的文件类型: {ext}")

    if not full_text or not full_text.strip():
        raise ParseError("文档解析后内容为空，无法提取有效文本")

    return ParsedDocument(
        doc_id=doc_id,
        filename=filename,
        file_type=file_type,
        full_text=full_text,
        text_chunks=[],   # filled by chunker
        images=images,     # captions filled later
    )
